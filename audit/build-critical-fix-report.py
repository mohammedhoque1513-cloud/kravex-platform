from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"C:\Users\emdad\Projects\KRAVEX")
OUT = ROOT / "audit" / "KRAVEX-critical-fix-pass-report-2026-06-14.docx"


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    props.append(node)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(46, 116, 181)
    return p


def status_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, label in enumerate(["Check", "Result", "Evidence", "Remaining action"]):
        shade(table.rows[0].cells[index], "F2F4F7")
        run = table.rows[0].cells[index].paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(9)
    for check, result, evidence, action in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, [check, result, evidence, action]):
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
        cells[1].paragraphs[0].runs[0].bold = True
        shade(cells[1], "E8F5E9" if result == "PASS" else "FFF4CC" if result == "BLOCKED" else "FDECEC")
    return table


def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def image(doc, path, caption, width=5.8):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(85, 85, 85)


doc = Document()
section = doc.sections[0]
section.top_margin = section.bottom_margin = Inches(1)
section.left_margin = section.right_margin = Inches(1)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)
doc.styles["Normal"].paragraph_format.space_after = Pt(6)

title = doc.add_paragraph()
run = title.add_run("KRAVEX Critical Fix Pass")
run.bold = True
run.font.size = Pt(24)
doc.add_paragraph(
    f"Evidence-based re-verification following the 07/06/2026 critical report. "
    f"Generated {datetime.now().strftime('%d/%m/%Y %H:%M')}."
)

heading(doc, "Executive verdict")
doc.add_paragraph(
    "The locally testable critical failures are fixed and proven against a production Next.js server "
    "and real PostgreSQL 16 database. Stripe test payments and hosted GitHub CI remain blocked by "
    "missing external credentials or an unpushed branch, so this pass is not represented as fully complete."
)

status_table(doc, [
    ("Environment stability", "PASS", "Canonical copy created at C:\\Users\\emdad\\Projects\\KRAVEX. Path does not contain OneDrive. Production build/start used.", "Delete the old OneDrive copy after Codex releases its workspace lock."),
    ("Production build", "PASS", "Next.js 15.5.19 optimized build completed successfully with 62 generated pages.", "None."),
    ("Login reload stability", "PASS", "10/10 consecutive /login requests returned HTTP 200.", "None."),
    ("Login credential leak", "PASS", "Raw HTML contains method=\"post\" and action=\"/api/auth/noop\"; SSR submit button is disabled until mounted.", "None."),
    ("Admin login", "PASS", "Playwright redirected real seeded admin to /admin/dashboard; wrong password returned plain-English error.", "None."),
    ("Client login", "PASS", "Playwright redirected seeded client to /client/dashboard.", "None."),
    ("PostgreSQL", "PASS", "Health endpoint returned db=connected; 4 migrations deployed and seed counts verified.", "Replace local DB URL with Railway for deployment."),
    ("Lead form", "PASS", "Valid POST returned 200 and persisted a LeadForm row; invalid POST returned 400 with field errors.", "Configure Resend/reCAPTCHA for production delivery and bot scoring."),
    ("Authenticated APIs", "PASS", "Admin and client API tests passed; client role was denied access to admin APIs.", "None."),
    ("Admin portal pages", "PASS", "11 authenticated admin pages returned 200 and remained under /admin.", "None."),
    ("Client portal pages", "PASS", "5 authenticated client pages returned 200 and remained under /client.", "None."),
    ("Public routes", "PASS", "16 production smoke routes returned HTTP 200.", "None."),
    ("Dependency security", "PASS", "audit-ci --high passed: 0 high, 0 critical. Four moderate advisories remain documented.", "Monitor Next/PostCSS and NextAuth/UUID upstream fixes."),
    ("Stripe payment flow", "BLOCKED", "Stripe keys and webhook secret are placeholders; no real test PaymentIntent/webhook can be proven.", "Provide Stripe test keys and run Stripe CLI forwarding."),
    ("Hosted CI", "BLOCKED", ".github/workflows/ci.yml exists, but the branch has not been pushed to GitHub.", "Push branch and verify the GitHub Actions run is green."),
])

heading(doc, "Verified database records")
for text in [
    "3 users, including one admin and two client portal users.",
    "2 clients: Patel Dental and Metro Roofing.",
    "3 prospects, 8 leads, 2 invoices, 2 campaigns and 3 settings.",
    "5 money vaults using the 30/5/20/10/35 allocation policy.",
    "LeadForm rows created by regression tests are persisted in PostgreSQL.",
]:
    bullet(doc, text)

heading(doc, "Automated test result")
doc.add_paragraph("Playwright critical suite: 10 passed, 0 failed.")
for text in [
    "Pre-hydration login safety and raw HTML method/action.",
    "Admin success login and wrong-password response.",
    "Client success login.",
    "Valid and invalid lead-form API submissions with database persistence.",
    "Admin protected API access with seeded data.",
    "Client API access and admin-route denial.",
    "Every admin page and every client page loaded inside the authenticated portal.",
]:
    bullet(doc, text)

heading(doc, "Evidence screenshots")
image(doc, ROOT / "audit" / "fix-evidence" / "admin" / "dashboard.png", "Authenticated admin dashboard after the fix.", 6.0)
image(doc, ROOT / "audit" / "fix-evidence" / "admin" / "payments.png", "Authenticated admin payments page after the fix.", 6.0)
image(doc, ROOT / "audit" / "fix-evidence" / "client" / "dashboard.png", "Authenticated client dashboard after the fix.", 6.0)
image(doc, ROOT / "audit" / "fix-evidence" / "client" / "leads.png", "Authenticated client leads page after the fix.", 6.0)
image(doc, ROOT / "audit" / "fix-evidence" / "client" / "account.png", "Authenticated client account page after the fix.", 6.0)

heading(doc, "Remaining blockers")
for text in [
    "Stripe: test secret, publishable key and webhook signing secret are not configured. Payment success, decline, receipt email and webhook vault allocation cannot be honestly marked PASS.",
    "Resend: API key is not configured, so email delivery is skipped by design. Core lead persistence no longer depends on email.",
    "GitHub Actions: workflow is created locally but cannot be green until the branch is pushed to a GitHub repository.",
    "Railway/Vercel: deployment credentials are not available in this workspace. Local PostgreSQL is real and proven, but hosted persistence is not yet proven.",
    "The old OneDrive folder remains because the Codex desktop holds it open. The tested canonical copy is the non-OneDrive path above.",
]:
    bullet(doc, text)

heading(doc, "Security notes")
doc.add_paragraph(
    "Next.js was upgraded from 14.2.35 to 15.5.19. The original high-severity Next findings are gone. "
    "The remaining moderate advisories are a PostCSS version bundled inside Next and UUID bundled inside "
    "NextAuth 4.24.14. npm currently proposes unsafe framework downgrades rather than a valid fix; the CI "
    "gate blocks high and critical findings."
)

doc.save(OUT)
print(OUT)
