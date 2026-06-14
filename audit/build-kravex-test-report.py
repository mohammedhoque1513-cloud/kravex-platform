from __future__ import annotations

import json
import os
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"C:\Users\emdad\OneDrive\Documents\New project")
ASSETS = ROOT / "audit" / "kravex-report-assets"
SHOTS = ASSETS / "screenshots"
OUT = ROOT / "audit" / "KRAVEX-full-critical-testing-report.docx"


def load_json(name: str, default):
    path = ASSETS / name
    if not path.exists():
        return default
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


public = load_json("audit-results-public.json", {})
forms = load_json("audit-results-forms.json", {})
admin = load_json("audit-results-admin.json", {})


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(9)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(46, 116, 181) if level <= 2 else RGBColor(31, 77, 120)
    return p


def add_status_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Area", "Verdict", "Evidence", "Action"]
    for i, h in enumerate(headers):
        shade_cell(table.rows[0].cells[i], "F2F4F7")
        set_cell_text(table.rows[0].cells[i], h, bold=True)
    for area, verdict, evidence, action, fill in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], area, bold=True)
        set_cell_text(cells[1], verdict, bold=True, color="008A00" if verdict.startswith("PASS") else "9B1C1C" if verdict.startswith("FAIL") else "7A5A00")
        set_cell_text(cells[2], evidence)
        set_cell_text(cells[3], action)
        shade_cell(cells[1], fill)
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_image(doc, file_name, caption, width=5.8):
    path = SHOTS / file_name
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(85, 85, 85)


def count_public():
    routes = public.get("public", [])
    total = len(routes)
    passed = sum(1 for r in routes if r.get("ok"))
    mobile_total = sum(1 for r in public.get("mobile", []))
    mobile_pass = sum(1 for r in public.get("mobile", []) if r.get("ok"))
    return total, passed, mobile_total, mobile_pass


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.1

    title = doc.add_paragraph()
    title_run = title.add_run("KRAVEX Website & Platform Critical Test Report")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(10, 10, 10)
    subtitle = doc.add_paragraph()
    subtitle.add_run(f"Evidence-based audit of public website, admin portal, client portal, APIs, security, mobile rendering and reliability. Generated {datetime.now().strftime('%d/%m/%Y %H:%M')}.").italic = True

    add_heading(doc, "Executive verdict", 1)
    p = doc.add_paragraph()
    p.add_run("No BS result: ").bold = True
    p.add_run("the public website mostly loads and looks credible, but the platform is not production-ready. Authentication, lead submission, dependency security, and dev-server reliability have hard failures.")

    public_total, public_passed, mobile_total, mobile_pass = count_public()
    add_status_table(doc, [
        ("Public website pages", f"PASS {public_passed}/{public_total}", "All listed public pages returned 200 on desktop, tablet and phone in the automated route audit.", "Keep polishing copy/content, but loading is working.", "E8F5E9"),
        ("Mobile/tablet rendering", f"PASS {mobile_pass}/{mobile_total}", "No page-level horizontal overflow recorded in the public mobile/tablet probes.", "Continue manual visual QA on real devices.", "E8F5E9"),
        ("Admin portal access", "FAIL", "Admin login either leaked credentials via early native GET submission or stuck on Signing in during hydrated test. Protected API checks then returned 401.", "Fix login form method/action, prevent native credential GET, add robust auth tests.", "FDECEC"),
        ("Client portal access", "PARTIAL", "Client dashboard was reached in one hydrated run and screenshot captured; earlier client route captures were login-page false positives.", "Retest after fixing auth flow; add stable E2E auth tests.", "FFF4CC"),
        ("Lead/contact forms", "FAIL", "Forms render, but valid public lead-form API submission returned HTTP 500.", "Fix /api/lead-form persistence/email fallback and add API regression test.", "FDECEC"),
        ("Security dependency audit", "FAIL", "npm audit reports 7 vulnerabilities: 4 high and 3 moderate. Includes Next.js and NextAuth related advisories.", "Upgrade framework/auth packages on a controlled branch and retest.", "FDECEC"),
        ("TypeScript", "PASS", "npx tsc --noEmit completed with exit code 0.", "Keep strict checks in CI.", "E8F5E9"),
        ("Local reliability", "FAIL", "Server reproduced .next chunk corruption: Cannot find module './1682.js'; clean .next restart restored /login.", "Move project outside OneDrive or run production build/start for serious testing; add health checks.", "FDECEC"),
    ])

    add_heading(doc, "Critical failures", 1)
    for item in [
        "Login can expose credentials in the URL if the form is submitted before React hydration. This was observed as /login?email=...&password=... in the browser URL. Passwords are intentionally redacted from this report.",
        "Hydrated admin login did not reach /admin/dashboard in the final captured run; screenshot shows the page stuck on Signing in.",
        "Public lead-form API accepted the request path but returned HTTP 500 for a valid test enquiry, so enquiry capture is not reliable.",
        "Admin API endpoints correctly reject unauthenticated requests with 401, but authenticated admin API access was not proven because admin login failed.",
        "The local dev server became unstable during testing and served 500 errors for Next chunks until .next was deleted and the server restarted.",
        "npm audit is failing. This is not safe to ignore for a money-handling/client-data platform.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Confirmed strengths", 1)
    for item in [
        "Public marketing routes exist and returned HTTP 200 across desktop, tablet and phone.",
        "The public homepage now explains the service clearly: UK lead generation for service businesses, fixed monthly retainer, lead delivery and client portal.",
        "Public pages have real page-specific content for How It Works, Services, Industries, Results, About, Contact and legal pages.",
        "Client dashboard can render a real portal-style view when client auth completes.",
        "Unauthenticated admin/client API calls return 401 instead of leaking data.",
        "The design system is no longer stylesheet-missing bland HTML during healthy server operation.",
        "TypeScript currently passes.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Route audit", 1)
    doc.add_paragraph(f"Public route audit: {public_passed}/{public_total} passed. Tested routes included /, /how-it-works, /services, /industries, /results, /about, /contact, /privacy-policy, /terms, /cookie-policy, /get-started, /signup, /login, /forgot-password, /reset-password and /activate across desktop/tablet/phone.")
    doc.add_paragraph("Admin route audit: route files exist, but final protected-route run failed because admin authentication did not complete. Earlier visual screenshots showed admin UI screens rendering, but the report verdict is based on the final protected access test.")
    doc.add_paragraph("Client route audit: client dashboard rendered in a hydrated login run. Other client sections require retest after auth hardening because earlier captures were polluted by login-page redirects.")

    add_heading(doc, "API and form audit", 1)
    add_status_table(doc, [
        ("Contact form UI", "PASS", "Form present on /contact with six inputs and Send Enquiry control.", "Test POST /api/contact after auth/server stability fixes.", "E8F5E9"),
        ("Lead form UI", "PASS", "Form present on /get-started with eleven inputs and Request a Free Strategy Call control.", "Keep.", "E8F5E9"),
        ("Lead form POST", "FAIL", "POST /api/lead-form returned 500 for valid test lead.", "Fix route error path and verify DB/local fallback.", "FDECEC"),
        ("Admin APIs unauthenticated", "PASS", "Admin API endpoints returned 401 when no valid admin session existed.", "Keep.", "E8F5E9"),
        ("Authenticated admin APIs", "FAIL/UNPROVEN", "Could not prove because admin login failed.", "Retest after login fix.", "FDECEC"),
        ("Stripe/payment APIs", "UNPROVEN", "Not exercised with real Stripe credentials/webhooks in this local run.", "Use Stripe test mode and webhook replay.", "FFF4CC"),
    ])

    add_heading(doc, "Screenshots", 1)
    add_image(doc, "public-home-phone.png", "Public homepage on phone: service proposition visible and styled.", width=3.0)
    add_image(doc, "public-services-phone.png", "Services page on phone: pricing/content page loaded.", width=3.0)
    add_image(doc, "public-industries-phone.png", "Industries page on phone: industry sections loaded.", width=3.0)
    add_image(doc, "admin-admin-dashboard-desktop.png", "Admin dashboard earlier visual capture: UI can render, but final auth run did not prove stable access.", width=6.0)
    add_image(doc, "auth-admin-after-login.png", "Final admin auth evidence: stuck on Signing in after credentials submitted.", width=5.6)
    add_image(doc, "auth-client-after-login.png", "Client auth evidence: client dashboard reached and rendered.", width=5.6)
    add_image(doc, "public-get-started-phone.png", "Lead form page on phone: form UI visible.", width=3.0)

    add_heading(doc, "Page-by-page verdict", 1)
    rows = []
    for route in public.get("public", []):
        if route.get("viewport") == "desktop":
            rows.append((route.get("label"), "PASS" if route.get("ok") else "FAIL", route.get("route"), "HTTP 200 and rendered content" if route.get("ok") else route.get("finalUrl"), "E8F5E9" if route.get("ok") else "FDECEC"))
    add_status_table(doc, rows[:16])

    add_heading(doc, "Priority fix list", 1)
    for item in [
        "Fix login immediately: set form method='post', block native GET credential submission, add autocomplete attributes, ensure NextAuth signIn always resolves, and add Playwright tests for admin/client login.",
        "Fix /api/lead-form 500. Log the actual exception, verify local fallback/PostgreSQL persistence, and prove valid submissions produce a saved LeadForm row or Prospect.",
        "Move serious testing away from OneDrive-backed .next or use production build/start. The current local dev setup can corrupt chunks under load.",
        "Upgrade vulnerable packages. Next.js 14.2.35 and NextAuth 4.24.14 are flagged by npm audit in this environment.",
        "Add E2E coverage for admin CRUD, client lead status updates, invoice creation, Stripe test payment success/failure, PDF generation and email send stubs.",
        "Connect PostgreSQL and run migrations/seed. Until then, this is a local-demo platform, not production data handling.",
        "Add CI gates: typecheck, build, npm audit, route smoke test, Playwright auth test, and lead-form POST test.",
    ]:
        add_number(doc, item)

    add_heading(doc, "Testing limits", 1)
    for item in [
        "No live PostgreSQL database, live Stripe, live Resend, Vercel Blob, Companies House or Cloudflare production account was connected during this test.",
        "Payment settlement, real payouts, digital vault withdrawals and bank deposits were not proven. Those require Stripe test/live configuration and reconciliation tests.",
        "The report is based on observed local behaviour on 07/06/2026. It does not certify production readiness.",
    ]:
        add_bullet(doc, item)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
