import json
import os
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
AUDIT = ROOT / "audit"
SCREENSHOTS = AUDIT / "screenshots"
OUT = AUDIT / "KRAVEX_Website_Test_Audit.docx"
IMG_OUT = AUDIT / "report_images"
IMG_OUT.mkdir(exist_ok=True)

BRAND_BLACK = "0A0A0A"
BRAND_GREY = "1A1A1A"
GOLD = "C9A84C"
GOLD_LIGHT = "E8C97A"
OFF_WHITE = "F5F5F0"
RED = "EF4444"
GREEN = "22C55E"
AMBER = "F59E0B"
MID = "D9D9D9"


def load_json(name):
    with open(AUDIT / name, "r", encoding="utf-8") as f:
        return json.load(f)


results = load_json("test-results.json")
interactions = load_json("interaction-results.json")


def money(n):
    return f"GBP {n / 100:,.2f}"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color="000000", size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def status_word(ok):
    return "PASS" if ok else "FAIL"


def status_fill(status):
    s = str(status).upper()
    if s in ("PASS", "OK", "WORKS", "200"):
        return "DDFBE8"
    if s in ("PARTIAL", "EXPECTED 401", "PROTECTED"):
        return "FEF3C7"
    if s in ("FAIL", "BROKEN", "500"):
        return "FEE2E2"
    return "EFEFEF"


def add_status_table(doc, headers, rows, widths=None, font_size=8.0):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i], BRAND_BLACK)
        set_cell_text(hdr[i], h, bold=True, color=OFF_WHITE, size=font_size)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            fill = "FFFFFF"
            if i == 0 and str(val).upper() in ("PASS", "FAIL", "PARTIAL", "PROTECTED", "EXPECTED 401"):
                fill = status_fill(str(val))
            shade_cell(cells[i], fill)
            set_cell_text(cells[i], val, bold=(i == 0), color="000000", size=font_size)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor.from_string(GOLD if level == 1 else BRAND_BLACK)
    return p


def add_callout(doc, title, body, fill="F7F3E5", border=GOLD):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(BRAND_BLACK)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.name = "Arial"
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor.from_string(BRAND_BLACK)
    doc.add_paragraph()


def prep_image(src: Path, max_width=1400):
    img = Image.open(src)
    img = img.convert("RGB")
    if img.width > max_width:
        ratio = max_width / img.width
        img = img.resize((max_width, int(img.height * ratio)))
    dest = IMG_OUT / (src.stem + ".jpg")
    img.save(dest, quality=78, optimize=True)
    return dest


def add_screenshot(doc, title, path, width=6.2, max_height=7.0):
    if not Path(path).exists():
        p = doc.add_paragraph(f"Screenshot missing: {path}")
        p.runs[0].font.color.rgb = RGBColor.from_string(RED)
        return
    prepared = prep_image(Path(path))
    with Image.open(prepared) as image:
        aspect = image.width / image.height
        fit_width = min(width, max_height * aspect)
    doc.add_picture(str(prepared), width=Inches(fit_width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(BRAND_BLACK)


def body_para(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(9.5)
    r.bold = bold
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.6)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(9.5)
for i in range(1, 4):
    styles[f"Heading {i}"].font.name = "Arial"
    styles[f"Heading {i}"].font.bold = True
styles["Heading 1"].font.size = Pt(18)
styles["Heading 2"].font.size = Pt(13)
styles["Heading 3"].font.size = Pt(10.5)

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("KRAVEX")
r.bold = True
r.font.name = "Arial"
r.font.size = Pt(36)
r.font.color.rgb = RGBColor.from_string(GOLD)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p2.add_run("Website and Platform Testing Report")
r.bold = True
r.font.name = "Arial"
r.font.size = Pt(22)
r.font.color.rgb = RGBColor.from_string(BRAND_BLACK)
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p3.add_run("Critical audit of public site, authentication, admin portal, client portal, API routes, payments, screenshots and verified behaviour.")
r.font.name = "Arial"
r.font.size = Pt(11)
r.font.color.rgb = RGBColor.from_string("555555")

meta_rows = [
    ["Test target", "http://localhost:3000"],
    ["Run date", "01/06/2026"],
    ["Tester", "Codex local browser/Puppeteer audit"],
    ["Evidence files", "audit/test-results.json, audit/interaction-results.json, audit/screenshots"],
    ["Verdict", "NOT PRODUCTION READY"],
]
add_status_table(doc, ["Item", "Value"], meta_rows, widths=[4.2, 12.5], font_size=9)
add_callout(
    doc,
    "Bottom line",
    "The app now has broad visual route coverage, but it is not a working production platform. Pages render, but live database-backed forms, real login, CRUD, payments, emails, reconciliation and backups are not proven operational. Public form APIs and scheduled job endpoints fail because the configured PostgreSQL database does not exist.",
    fill="FEE2E2",
)
doc.add_page_break()

add_heading(doc, "1. Executive Verdict", 1)
body_para(doc, "No sugar-coating: this is a visual platform shell with partial front-end interactions, not a finished production system.")

summary_rows = [
    ["PASS", "Page rendering", "34/34 tested routes returned 200 and captured screenshots after fixing the invoice-create page crash."],
    ["PASS", "Styling loaded", "CSS and JavaScript assets loaded in the route audit; the unstyled page issue was resolved by restarting after the code fix."],
    ["FAIL", "Database-backed function", "Public contact, lead form, reconciliation and backup endpoints fail with Prisma: Database `kravex` does not exist."],
    ["FAIL", "Real seeded login", "The seeded admin credentials did not redirect to the admin dashboard in the browser test."],
    ["PARTIAL", "Admin UI", "Admin pages render. Some controls work locally, but most data/actions are shells or protected APIs without verified persistence."],
    ["PARTIAL", "Client UI", "Client pages render. The leads page has filter UI but no real lead table/card data path verified."],
    ["FAIL", "Payments and vaults", "The money-split ledger code exists, but Stripe payment, payout, vault balances, emails and reconciliation were not end-to-end verified."],
]
add_status_table(doc, ["Result", "Area", "Certainty"], summary_rows, widths=[2.4, 4.0, 10.4], font_size=8.5)

add_heading(doc, "2. What Was Actually Tested", 1)
body_para(doc, "This report is based on live local testing against http://localhost:3000, Puppeteer page loads, screenshots, API requests and targeted interaction checks. It does not assume that code works because it exists.")
coverage_rows = [
    ["Pages", str(len(results["pages"])), "Loaded in browser, status checked, CSS/JS checked, screenshot captured."],
    ["API routes sampled", str(len(results["apis"])), "Direct HTTP calls made to public, admin, client, Stripe and job endpoints."],
    ["Interaction checks", str(len(interactions)), "Login, demo link, dashboard modal, leads search, settings tabs, invoice calculator, client shell checks."],
    ["Screenshots captured", str(len(list(SCREENSHOTS.glob('*.png')))), "Stored in audit/screenshots and embedded/linked in this report appendix."],
]
add_status_table(doc, ["Test group", "Count", "Method"], coverage_rows, widths=[4.2, 2.2, 10.4], font_size=8.5)

add_heading(doc, "3. Critical Blockers", 1)
blockers = [
    ["1", "PostgreSQL database missing", "Forms, jobs and any Prisma-backed function cannot operate. Server log confirms: Database `kravex` does not exist."],
    ["2", "Production login not proven", "The expected seeded admin login did not complete in the browser test. Demo navigation exists, but that is not secure authentication."],
    ["3", "APIs are not usable end-to-end", "Authenticated APIs correctly reject unauthenticated calls with 401, but there is no verified authenticated test path and public APIs return 500."],
    ["4", "Admin portal is not a full working management system", "Pages exist but major workflows such as create/edit/delete/convert/send/pay/export are not proven to persist to the database."],
    ["5", "Client portal is not a real client operating area yet", "Client views render but do not show verified client-specific live data, invoice payment, messaging or account update flows."],
    ["6", "Money handling cannot be trusted yet", "Vault split rules are coded, but there is no verified Stripe success webhook, balance reconciliation, bank payout, or failure recovery test."],
    ["7", "Email and PDF flow unproven", "Resend, invoice PDF storage and transactional emails were not verified because live credentials/storage are not configured."],
]
add_status_table(doc, ["Priority", "Blocker", "Evidence"], blockers, widths=[1.5, 5.2, 10.1], font_size=8)

add_heading(doc, "4. Page-by-Page Result", 1)
page_rows = []
for p in results["pages"]:
    result = "PASS" if p["status"] == 200 and p["loaded"] and p["styled"] and not p["error"] else "FAIL"
    page_rows.append([result, p["name"], p["url"], str(p["status"]), "Yes" if p["styled"] else "No", os.path.basename(p["screenshot"])])
add_status_table(doc, ["Result", "Page", "URL", "HTTP", "Styled", "Screenshot"], page_rows, widths=[2, 4.4, 4.2, 1.5, 1.6, 3.5], font_size=7)

add_heading(doc, "5. API Route Result", 1)
api_rows = []
for a in results["apis"]:
    if a["status"] == 401:
        verdict = "PROTECTED"
        evidence = "Rejected unauthenticated request: " + str(a.get("response", ""))[:80]
    elif a["ok"]:
        verdict = "PASS"
        evidence = str(a.get("response", ""))[:80]
    else:
        verdict = "FAIL"
        evidence = str(a.get("response") or a.get("error") or "No response body")[:120]
    api_rows.append([verdict, a["method"], a["url"], str(a["status"]), evidence])
add_status_table(doc, ["Verdict", "Method", "Route", "HTTP", "Evidence"], api_rows, widths=[2.2, 2.0, 5.2, 1.6, 5.8], font_size=7)
add_callout(
    doc,
    "API interpretation",
    "401 on admin/client routes is acceptable route protection only for unauthenticated requests. It is not proof that the feature works after login. 500 on public form and job routes is a hard failure.",
    fill="FEF3C7",
)

add_heading(doc, "6. Interaction Checks", 1)
interaction_rows = []
for item in interactions:
    interaction_rows.append([status_word(item["ok"]), item["name"], f"{item['durationMs']} ms", item.get("error", "") or "Observed by browser automation"])
add_status_table(doc, ["Result", "Check", "Time", "Evidence"], interaction_rows, widths=[2, 7.2, 2.2, 5.4], font_size=7.5)

add_heading(doc, "7. Strengths", 1)
strengths = [
    "The project has broad route coverage: public, auth, admin and client sections all have pages that load.",
    "The black/gold KRAVEX visual direction is present and CSS loads correctly after the invoice-page fix.",
    "The Prisma schema and newer money-vault ledger models give the right starting structure for real finance tracking.",
    "Admin settings tabs, leads search input and invoice VAT calculation showed working front-end behaviour.",
    "The dashboard now reflects zero real clients/money rather than pretending the business has live revenue.",
]
for s in strengths:
    body_para(doc, s)

add_heading(doc, "8. Flaws Requiring Work", 1)
flaws = [
    ["Database", "Create the PostgreSQL database, run Prisma migrations, seed safely, and prove CRUD works from the UI."],
    ["Authentication", "Make real NextAuth credentials login work with Jordan's admin user and client portal users. Remove/demo-gate shortcut before production."],
    ["Admin dashboard", "Replace static blocks with live database queries, real empty states, graph data and working quick actions."],
    ["Admin workflows", "Prospects, clients, leads, campaigns, invoices, payments, reports and settings need verified create/edit/delete/export/send flows."],
    ["Client portal", "Wire dashboard, leads, invoices, messages and account pages to authenticated client-specific records."],
    ["Payments", "End-to-end Stripe PaymentIntent, webhook, invoice paid state, vault ledger entries, receipt email and payout request must be tested with Stripe test mode."],
    ["Vaults", "Add visible ledger audit trail, reconciliation checks, locked tax reserve rules and withdrawal limits before using owner pay."],
    ["Email/PDF", "Configure Resend, Blob storage and Puppeteer PDF generation; test invoice email, receipt email, password reset and activation."],
    ["Testing", "Add automated tests for API auth, database transactions, payment webhook idempotency, form validation and role protection."],
]
add_status_table(doc, ["Area", "Required work"], flaws, widths=[3.6, 13.2], font_size=8.2)

add_heading(doc, "9. Money Flow Assessment", 1)
body_para(doc, "Requested flow tested conceptually against implemented routes and directly against available endpoints. Certainty: the complete flow does not work today.")
money_rows = [
    ["Client pays by card", "FAIL", "POST /api/stripe/pay-invoice returns 401 without authenticated client session. No successful Stripe test charge observed."],
    ["Webhook fires", "NOT VERIFIED", "Webhook route exists, but no signed Stripe event was tested end-to-end."],
    ["30/5/20/10/35 vault split", "PARTIAL", "Split policy is coded in lib/money.ts, but no real payment created ledger entries in the database."],
    ["Invoice marked paid", "FAIL", "No successful payment/webhook test; invoice state change not proven."],
    ["Client receipt email", "FAIL", "Resend not configured/tested."],
    ["Jordan notification", "FAIL", "Resend not configured/tested."],
    ["Nightly reconciliation", "FAIL", "POST /api/jobs/reconcile returns 500 because database `kravex` does not exist."],
    ["Midnight backup", "FAIL", "POST /api/jobs/backup returns 500 because database `kravex` does not exist."],
    ["Owner deposit", "NOT READY", "Withdrawal route exists but is not verified and should only request Stripe payouts to a verified bank account, not store/hold money internally."],
]
add_status_table(doc, ["Step", "Verdict", "Evidence"], money_rows, widths=[4.8, 2.8, 9.2], font_size=8)

add_heading(doc, "10. Screenshot Evidence", 1)
body_para(doc, "Every tested route produced a screenshot. Representative full-width screenshots are below; the appendix includes the full screenshot inventory.")
key_shots = [
    ("Public home page", SCREENSHOTS / "home.png"),
    ("Admin dashboard", SCREENSHOTS / "admin_dashboard.png"),
    ("Admin payments and vault section", SCREENSHOTS / "admin_payments.png"),
    ("Invoice creator", SCREENSHOTS / "admin_invoices_new.png"),
    ("Client dashboard", SCREENSHOTS / "client_dashboard.png"),
    ("Login page after failed credential test", SCREENSHOTS / "interaction-login-admin.png"),
]
for title, shot in key_shots:
    add_screenshot(doc, title, shot, width=6.2)

doc.add_page_break()
add_heading(doc, "Appendix A: Full Screenshot Inventory", 1)
body_para(doc, "These are the captured snapshots used as evidence. They are included here so the visual state of every tested page can be inspected from the document.")
for p in results["pages"]:
    add_screenshot(doc, f"{p['name']} - {p['url']}", p["screenshot"], width=5.9)
    if p["name"] in ("Login", "Admin Dashboard", "Admin Payments", "Client Dashboard"):
        body_para(doc, "Audit note: This page loaded visually, but visual load alone is not proof of backend functionality.")

add_heading(doc, "Appendix B: Raw Failure Evidence", 1)
server_evidence = [
    "Prisma public form failure: Invalid prisma.leadForm.create() invocation. Database `kravex` does not exist.",
    "Prisma job failure: Invalid prisma.reconciliationRun.create() invocation. Database `kravex` does not exist.",
    "Prisma job failure: Invalid prisma.backupRun.create() invocation. Database `kravex` does not exist.",
    "Public endpoints POST /api/contact and POST /api/lead-form returned HTTP 500.",
    "Authenticated API samples returned HTTP 401 without a signed-in session.",
]
for line in server_evidence:
    body_para(doc, line)

add_heading(doc, "Final Decision", 1)
add_callout(
    doc,
    "Certainty statement",
    "If the question is 'does the platform work as a real production KRAVEX business system today?', the answer is no. If the question is 'does the current build show the rough pages and brand shell?', the answer is yes. The next real milestone is not more decoration; it is database setup, real authentication, and one complete tested workflow from contact form to client payment to vault ledger.",
    fill="FEE2E2",
)

doc.save(OUT)
print(str(OUT))
